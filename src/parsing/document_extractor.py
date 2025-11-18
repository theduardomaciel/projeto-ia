"""Utilidades para extrair texto de documentos suportados.

Atualmente oferece suporte a arquivos .txt, .pdf e .docx, com
tratamento consistente de logs e normalizacoes simples.
"""

from __future__ import annotations

from pathlib import Path
from typing import Any, Callable

TextReader = Callable[[Path], str]


class DocumentExtractor:
    """Extrai texto bruto de diferentes formatos de curriculo."""

    SUPPORTED_EXTENSIONS = frozenset({".txt", ".pdf", ".docx"})

    def __init__(
        self,
        text_reader: TextReader | None = None,
        logger: Callable[[str, str], None] | None = None,
    ) -> None:
        self._text_reader = text_reader or self._default_text_reader
        self._logger = logger

    @property
    def supported_extensions(self) -> frozenset[str]:
        return self.SUPPORTED_EXTENSIONS

    def extract_text(self, path: str | Path) -> str:
        file_path = Path(path)
        if not file_path.exists():
            raise FileNotFoundError(f"Arquivo nao encontrado: {file_path}")

        suffix = file_path.suffix.lower()
        if suffix == ".txt":
            text = self._text_reader(file_path)
        elif suffix == ".pdf":
            text = self._extract_pdf(file_path)
        elif suffix == ".docx":
            text = self._extract_docx(file_path)
        else:
            raise ValueError(
                f"Extensao nao suportada '{suffix or '<sem extensao>'}'. "
                "Formatos aceitos: .txt, .pdf, .docx"
            )

        cleaned = self._post_process(text)
        if self._logger:
            self._logger(
                "document_extracted",
                f"path={file_path} ext={suffix or 'unknown'} chars={len(cleaned)}",
            )
        return cleaned

    # ------------------------------------------------------------------
    # Metodos auxiliares
    # ------------------------------------------------------------------
    def _default_text_reader(self, path: Path) -> str:
        return path.read_text(encoding="utf-8", errors="ignore")

    def _post_process(self, text: str) -> str:
        if not text:
            return ""
        normalized = text.replace("\r\n", "\n").replace("\r", "\n")
        return normalized.strip()

    def _extract_pdf(self, path: Path) -> str:
        try:
            import pdfplumber  # type: ignore
        except ImportError as exc:  # pragma: no cover - dependencia externa
            raise RuntimeError(
                "Dependencia 'pdfplumber' e obrigatoria para extrair PDFs. "
                "Execute 'pip install pdfplumber'."
            ) from exc

        chunks: list[str] = []
        with pdfplumber.open(str(path)) as pdf:  # type: ignore[arg-type]
            for page in pdf.pages:
                page_text = (page.extract_text() or "").replace("\u00a0", " ").strip()
                if page_text:
                    chunks.append(page_text)
        return "\n".join(chunks)

    def _extract_docx(self, path: Path) -> str:
        try:
            from docx import Document  # type: ignore
        except ImportError as exc:  # pragma: no cover - dependencia externa
            raise RuntimeError(
                "Dependencia 'python-docx' e obrigatoria para extrair DOCX. "
                "Execute 'pip install python-docx'."
            ) from exc

        document = Document(str(path))
        paragraphs = [p.text.strip() for p in document.paragraphs if p.text.strip()]
        table_entries = self._extract_docx_tables(document)
        return "\n".join(paragraphs + table_entries)

    def _extract_docx_tables(self, document: Any) -> list[str]:
        table_chunks: list[str] = []
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    table_chunks.append(" ".join(cells))
        return table_chunks
