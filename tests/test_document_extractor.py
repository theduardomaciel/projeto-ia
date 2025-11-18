from __future__ import annotations

import sys
from pathlib import Path

import pytest

from src.parsing.document_extractor import DocumentExtractor


def test_extract_text_from_plain_file(tmp_path: Path) -> None:
    file_path = tmp_path / "sample.txt"
    file_path.write_text("Primeira linha\nSegunda linha", encoding="utf-8")

    extractor = DocumentExtractor()
    text = extractor.extract_text(file_path)

    assert "Primeira linha" in text
    assert "Segunda linha" in text


def test_extract_text_from_docx(tmp_path: Path) -> None:
    pytest.importorskip("docx")
    from docx import Document  # type: ignore

    doc_path = tmp_path / "sample.docx"
    document = Document()
    document.add_paragraph("Resumo profissional")
    table = document.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Python"
    table.rows[0].cells[1].text = "LLMs"
    document.save(doc_path)

    extractor = DocumentExtractor()
    text = extractor.extract_text(doc_path)

    assert "Resumo profissional" in text
    assert "Python LLMs" in text


def test_extract_text_from_pdf_with_stub(monkeypatch, tmp_path: Path) -> None:
    pdf_path = tmp_path / "sample.pdf"
    pdf_path.write_bytes(b"%PDF-1.4\n%\xe2\xe3\xcf\xd3\n")

    class DummyPage:
        def __init__(self, text: str) -> None:
            self._text = text

        def extract_text(self) -> str:
            return self._text

    class DummyPDF:
        def __init__(self) -> None:
            self.pages = [DummyPage("Pagina 1"), DummyPage("Pagina 2")]

        def __enter__(self) -> "DummyPDF":  # pragma: no cover - helper
            return self

        def __exit__(self, *_) -> None:  # pragma: no cover - helper
            return None

    class DummyModule:
        def open(self, path) -> DummyPDF:  # pragma: no cover - helper
            assert Path(path) == pdf_path
            return DummyPDF()

    monkeypatch.setitem(sys.modules, "pdfplumber", DummyModule())

    extractor = DocumentExtractor()
    text = extractor.extract_text(pdf_path)

    assert "Pagina 1" in text
    assert "Pagina 2" in text
